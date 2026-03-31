import io
import os
from pathlib import Path

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF, handling multi-column and table layouts."""
    if not PDF_AVAILABLE:
        raise RuntimeError("pdfplumber is not installed.")
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            # Try words-based extraction for better layout handling
            words = page.extract_words(x_tolerance=3, y_tolerance=3)
            if words:
                # Sort by vertical position then horizontal (handles multi-column)
                line_map: dict = {}
                for w in words:
                    y_key = round(w["top"] / 5) * 5  # Group within 5pt bands
                    line_map.setdefault(y_key, []).append(w)
                for y_key in sorted(line_map.keys()):
                    line_words = sorted(line_map[y_key], key=lambda w: w["x0"])
                    text_parts.append(" ".join(w["text"] for w in line_words))
            else:
                # Fallback to simple extraction
                raw = page.extract_text()
                if raw:
                    text_parts.append(raw)
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX files."""
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not installed.")
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Also extract table text
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_text_from_image(file_bytes: bytes) -> str:
    """Extract text from images using OCR (requires Tesseract)."""
    if not OCR_AVAILABLE:
        raise RuntimeError(
            "pytesseract/Pillow not installed, or Tesseract OCR binary not found. "
            "Install Tesseract from https://github.com/UB-Mannheim/tesseract/wiki"
        )
    image = Image.open(io.BytesIO(file_bytes))
    # Convert to RGB if needed
    if image.mode not in ("RGB", "L"):
        image = image.convert("RGB")
    text = pytesseract.image_to_string(image, config="--psm 6")
    return text


def parse_file(file_bytes: bytes, filename: str) -> str:
    """Dispatch file parsing based on extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff"):
        return extract_text_from_image(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
